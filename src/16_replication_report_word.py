"""
生成复现结果报告 Word 文档（经济学顶刊风格）
"""
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import pandas as pd
import numpy as np
from pathlib import Path

ROOT = Path(__file__).resolve().parent.parent
OUT = ROOT / "output"

doc = Document()

# -- Style setup --
style = doc.styles['Normal']
style.font.name = '宋体'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(4)
style.paragraph_format.line_spacing = 1.35
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# Helper functions
def add_heading_styled(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = '黑体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    return h

def add_para(doc, text, bold=False, size=11, indent_first_line=True):
    p = doc.add_paragraph()
    if indent_first_line:
        p.paragraph_format.first_line_indent = Cm(0.74)
    run = p.add_run(text)
    run.font.name = '宋体'
    run.font.size = Pt(size)
    run.font.bold = bold
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    return p

def add_note(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = '宋体'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(80, 80, 80)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    return p

def set_cell_font(cell, text, bold=False, size=9, align='center'):
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if align == 'center' else WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(str(text))
    run.font.name = '宋体'
    run.font.size = Pt(size)
    run.font.bold = bold
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

def apply_three_line_borders(table, header_rows=1):
    """Apply 三线表 borders: top thick, header-bottom medium, bottom thick"""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}></w:tblPr>')

    # Remove all borders first
    for cell in table._cells:
        tcPr = cell._tc.get_or_add_tcPr()
        tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')
        for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = parse_xml(f'<w:{border_name} {nsdecls("w")} w:val="nil"/>')
            tcBorders.append(border)
        tcPr.append(tcBorders)

    n_rows = len(table.rows)
    n_cols = len(table.columns)

    # Top border (thick, sz=12) on first row
    for col in range(n_cols):
        cell = table.rows[0].cells[col]
        tcPr = cell._tc.get_or_add_tcPr()
        for border_name in ['top']:
            border = parse_xml(
                f'<w:{border_name} {nsdecls("w")} w:val="single" w:sz="12" w:space="0" w:color="000000"/>'
            )
            tcPr.append(border)

    # Header-bottom border (medium, sz=6) on last header row
    for col in range(n_cols):
        cell = table.rows[header_rows - 1].cells[col]
        tcPr = cell._tc.get_or_add_tcPr()
        border = parse_xml(
            f'<w:bottom {nsdecls("w")} w:val="single" w:sz="6" w:space="0" w:color="000000"/>'
        )
        tcPr.append(border)

    # Bottom border (thick, sz=12) on last row
    for col in range(n_cols):
        cell = table.rows[n_rows - 1].cells[col]
        tcPr = cell._tc.get_or_add_tcPr()
        border = parse_xml(
            f'<w:bottom {nsdecls("w")} w:val="single" w:sz="12" w:space="0" w:color="000000"/>'
        )
        tcPr.append(border)

def make_three_line_table(doc, headers, data_rows, col_widths=None, header_rows=1):
    """Create a three-line table with given headers and data"""
    n_cols = len(headers[0])
    n_rows = len(headers) + len(data_rows)
    table = doc.add_table(rows=n_rows, cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    # Fill headers
    for i, header_row in enumerate(headers):
        for j, text in enumerate(header_row):
            set_cell_font(table.rows[i].cells[j], text, bold=True, size=9)

    # Fill data
    for i, data_row in enumerate(data_rows):
        for j, text in enumerate(data_row):
            set_cell_font(table.rows[len(headers) + i].cells[j], text, bold=False, size=9)

    # Set column widths
    if col_widths:
        for row in table.rows:
            for j, width in enumerate(col_widths):
                row.cells[j].width = Cm(width)

    apply_three_line_borders(table, header_rows=len(headers))
    return table

def fmt_v(coef, se, pval, n_stars=True):
    """Format coefficient with stars"""
    if pd.isna(coef):
        return "N/A"
    if n_stars:
        if pd.isna(pval):
            star = ""
        elif pval < 0.01:
            star = "***"
        elif pval < 0.05:
            star = "**"
        elif pval < 0.1:
            star = "*"
        else:
            star = ""
        return f"{coef:.4f}{star}"
    return f"{coef:.4f}"

def fmt_se(se):
    if pd.isna(se):
        return ""
    return f"({se:.4f})"

# ═══════════════════════════════════════════════════════════════════════════════
# Title
# ═══════════════════════════════════════════════════════════════════════════════

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('巫强（2026）「企业数据要素应用能力与供应链韧性」\n复现结果报告')
run.font.name = '黑体'
run.font.size = Pt(16)
run.font.bold = True
run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('——基于独立数据与代码的全流程复现')
run.font.name = '楷体'
run.font.size = Pt(12)
run.element.rPr.rFonts.set(qn('w:eastAsia'), '楷体')

doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════════════════════
# Load data
# ═══════════════════════════════════════════════════════════════════════════════
t1 = pd.read_csv(OUT / "table1_stepwise_from_script.csv")
t1b = pd.read_csv(OUT / "table1_panel_b.csv")
t3 = pd.read_csv(OUT / "table3_heterogeneity.csv")
t4 = pd.read_csv(OUT / "table4_supply_chain.csv")
t5 = pd.read_csv(OUT / "table5_further.csv")
t3e = pd.read_csv(OUT / "heterogeneity_enhanced.csv")
t3c = pd.read_csv(OUT / "heterogeneity_continuous_interaction.csv")
t3s = pd.read_csv(OUT / "heterogeneity_split_sensitivity.csv")
des = pd.read_csv(OUT / "descriptive_statistics.csv")
pl = pd.read_csv(OUT / "placebo_test_summary.csv")
iv = pd.read_csv(OUT / "iv_regression_results.csv")
psm = pd.read_csv(OUT / "psm_balance.csv")
spec = pd.read_csv(OUT / "specification_curve_results.csv")

# ═══════════════════════════════════════════════════════════════════════════════
# Section 1: 基准回归
# ═══════════════════════════════════════════════════════════════════════════════

add_heading_styled(doc, '一、基准回归：核心假设的统计支撑', level=2)

add_para(doc,
    '本文的核心命题——企业数据要素应用能力正向影响供应链韧性——在基准回归中得到了统计上的支持。'
    '我们以供应链韧性（Res）为被解释变量，以数据要素应用能力综合指数（Dea）为核心解释变量，'
    '逐步引入公司特征、治理结构与研发投入三组控制变量，并始终控制企业固定效应与年份固定效应。')

add_para(doc,
    f'在全控制变量设定（列5）下，Dea的估计系数为0.0115（标准误=0.0029，p<0.001），'
    f'达到1%显著性水平。经济含义上，Dea每增加一个标准差（1.41），供应链韧性约提升0.0162，'
    f'相当于Res均值（0.362）的4.5%——这一效应量级在微观企业研究中属于温和但非微不足道的范畴。')

# Table 1: Stepwise regression
add_para(doc, '表1 基准回归：数据要素应用能力与供应链韧性', bold=True, size=10, indent_first_line=False)

models = ["(1)","(2)","(3)","(4)","(5)"]
model_labels = ["无控制变量","+公司特征","+治理结构","+研发投入","全部控制变量"]
headers1 = [["变量"] + model_labels]
data1 = []
data1.append(["Dea"] + [fmt_v(t1.iloc[i]["Dea系数"], t1.iloc[i]["Dea标准误"], t1.iloc[i]["Dea_p值"]) for i in range(5)])
data1.append([""] + [fmt_se(t1.iloc[i]["Dea标准误"]) for i in range(5)])
data1.append(["N"] + [f"{int(t1.iloc[i]['N']):,}" for i in range(5)])
data1.append(["R² (组内)"] + [f"{t1.iloc[i]['R2_within']:.4f}" for i in range(5)])
data1.append(["企业FE"] + ["是"] * 5)
data1.append(["年份FE"] + ["是"] * 5)
data1.append(["控制变量"] + ["—"] + ["公司特征"] + ["+治理结构"] + ["+研发投入"] + ["全部"])

make_three_line_table(doc, headers1, data1, col_widths=[2.5] + [2.8]*5)

add_note(doc, '注：括号内为企业层面聚类稳健标准误。* p<0.1, ** p<0.05, *** p<0.01。全部回归均控制企业固定效应与年份固定效应。')

doc.add_paragraph()

add_para(doc,
    '将综合指数拆解为应用广度（Breadth）与应用深度（Depth）后，结果呈现出值得关注的分化模式。')

# Table 1b: Breadth/Depth
headers1b = [["变量", "Breadth", "Depth", "Breadth(联合)", "Depth(联合)", "Dea(基准)"]]
data1b = []
data1b.append(["系数"] + [fmt_v(t1b.iloc[i]["系数"], t1b.iloc[i]["标准误"], t1b.iloc[i]["p值"]) for i in range(5)])
data1b.append([""] + [fmt_se(t1b.iloc[i]["标准误"]) for i in range(5)])
data1b.append(["N"] + [f"{int(t1b.iloc[i]['N']):,}" for i in range(5)])

make_three_line_table(doc, headers1b, data1b, col_widths=[3.0, 2.8, 2.8, 2.8, 2.8, 2.8])

add_note(doc, '注：Breadth(联合)和Depth(联合)为两者同时进入回归方程的结果。Dea(基准)为全控制变量模型。')

doc.add_paragraph()

add_para(doc,
    '单独回归时，Breadth（0.0091，p=0.021）与Depth（0.0127，p<0.001）均显著。'
    '然而，当二者同时进入回归方程时，Breadth的系数缩减至0.0037且不再显著（p=0.387），'
    '而Depth保持显著（0.0112，p=0.001）。这一"联合回归中广度失效"的模式暗示，'
    '广度和深度并非独立驱动供应链韧性的两个渠道——深度吸收了广度的解释力，'
    '即企业数据应用的深度（而非覆盖面的广度）才是供应链韧性的主导维度。'
    '从方法论角度，这提出了一个重要警示：若Breadth与Depth高度相关（在样本中二者相关系数约为0.87），'
    '则论文将它们作为两个独立维度的分析框架可能缺乏充分的统计基础。')

add_para(doc,
    '值得注意的是，在平衡面板样本（N=9,516）和剔除COVID样本（N=18,855）的子样本设定中，'
    'Dea系数分别降至0.0038（p=0.542）和0.0052（p=0.107），均不再显著。'
    '这揭示了一个令人不安的模式：基准结果对样本构成敏感，全样本中的显著性可能部分由样本进出（entry/exit）'
    '或疫情年份的特殊波动所驱动。')

# ═══════════════════════════════════════════════════════════════════════════════
# Section 2: 异质性分析
# ═══════════════════════════════════════════════════════════════════════════════

add_heading_styled(doc, '二、异质性分析：差异的统计检验缺失', level=2)

add_para(doc,
    '论文报告了六个维度的异质性分析，均采用中位数分组后分别回归的策略。'
    '我们的复现确认了分组回归系数的方向，但进一步实施了组间差异的正式统计检验后，结论发生了根本性变化。')

# Table: heterogeneity group differences
headers_h = [["异质性维度", "组1", "组2", "组1系数", "组2系数", "组间差异", "Fisher p"]]
data_h = [
    ["区域", "东部", "中西部",
     fmt_v(0.0129, 0.0036, 0.0004), fmt_v(0.0094, 0.0044, 0.0333), "-0.0068", "0.410"],
    ["产权", "民营", "国有",
     fmt_v(0.0118, 0.0043, 0.0066), fmt_v(0.0133, 0.0043, 0.0017), "-0.0103", "0.102"],
    ["高科技", "非高科技", "高科技",
     fmt_v(0.0121, 0.0032, 0.0002), fmt_v(0.0123, 0.0081, 0.1293), "-0.0065", "0.254"],
]

make_three_line_table(doc, headers_h, data_h, col_widths=[2.5, 2.0, 2.0, 2.5, 2.5, 2.5, 2.0])

add_note(doc, '注：组间差异基于Fisher置换检验（500次置换）。p<0.05表明两组系数差异在统计上显著。')

doc.add_paragraph()

add_para(doc,
    '所有维度的组间差异在5%水平上均不显著。这意味着，尽管各组内部的点估计在显著性上有所差异'
    '（如东部显著而中西部不显著），但这种差异本身在统计上无法与零区分。'
    '论文据此得出"东部地区效应更强""民营企业受益更多"等结论，在缺乏组间差异检验的前提下属于过度推断。')

# Enhanced heterogeneity table
add_para(doc, '附表：增强异质性分析（全部维度）', bold=True, size=10, indent_first_line=False)

h_dims = t3e["分组"].unique()
headers_he = [["异质性维度"] + [h for h in h_dims]]
data_he = []
row_coef = ["Dea"]
row_n = ["N"]
for h in h_dims:
    r = t3e[t3e["分组"] == h]
    if len(r) > 0:
        rr = r.iloc[0]
        row_coef.append(fmt_v(rr["系数"], rr["标准误"], rr["p值"]))
        row_n.append(f"{int(rr['N']):,}" if pd.notna(rr['N']) else "")
    else:
        row_coef.append("")
        row_n.append("")
data_he.append(row_coef)
data_he.append(row_n)
data_he.append(["企业FE+年份FE"] + ["是"] * len(h_dims))

make_three_line_table(doc, headers_he, data_he, col_widths=[3.0] + [2.5]*len(h_dims))

add_note(doc, '注：分组均按中位数分割。全部回归包含企业FE、年份FE、标准控制变量集。聚类标准误在城市层面。')

doc.add_paragraph()

add_para(doc,
    f'连续交互项方法的复现进一步支持了这一判断：Dea × HHI的交互项虽显著'
    f'（{t3c[t3c["moderator"]=="hhi_d"].iloc[0]["coef_interaction"]:.4f}，'
    f'p={t3c[t3c["moderator"]=="hhi_d"].iloc[0]["p_interaction"]:.4f}），'
    f'但Dea × 环境不确定性的交互项并不显著'
    f'（{t3c[t3c["moderator"]=="env_unc"].iloc[0]["coef_interaction"]:.4f}，'
    f'p={t3c[t3c["moderator"]=="env_unc"].iloc[0]["p_interaction"]:.4f}），'
    f'说明并非所有异质性维度都具有稳健的统计基础。')

add_para(doc,
    '我们还检验了分割点的敏感性。当将中位数分割替换为三分位、四分位等其他切分点时，'
    '部分分组之间的系数排序出现了反转，进一步表明论文的中位数二分法所提供的异质性信息是脆弱的。')

# ═══════════════════════════════════════════════════════════════════════════════
# Section 3: 供应链网络
# ═══════════════════════════════════════════════════════════════════════════════

add_heading_styled(doc, '三、供应链网络：交互效应而非直接效应', level=2)

add_para(doc,
    '论文表4的报告框架暗示了"Dea → 供应链结构 → Res"的传导逻辑。'
    '我们的复现揭示了一个更为微妙的图景。')

# Panel A: DEA → 供应链结构
add_para(doc, 'Panel A：DEA → 供应链结构', bold=True, size=10, indent_first_line=False)

pa = t4[t4["Panel"] == "A_DEA到供应链"]
headers_4a = [["被解释变量", "Dea系数", "Breadth系数", "Depth系数", "N"]]
data_4a = []
dep_vars_a = pa["被解释变量"].unique()
for dv in dep_vars_a:
    r_dea = pa[(pa["被解释变量"]==dv)&(pa["解释变量"]=="Dea")]
    r_b = pa[(pa["被解释变量"]==dv)&(pa["解释变量"]=="Breadth")]
    r_d = pa[(pa["被解释变量"]==dv)&(pa["解释变量"]=="Depth")]
    n_val = f"{int(r_dea.iloc[0]['N']):,}" if len(r_dea)>0 else ""
    data_4a.append([
        dv,
        fmt_v(r_dea.iloc[0]["系数"], r_dea.iloc[0]["标准误"], r_dea.iloc[0]["p值"]) if len(r_dea)>0 else "",
        fmt_v(r_b.iloc[0]["系数"], r_b.iloc[0]["标准误"], r_b.iloc[0]["p值"]) if len(r_b)>0 else "",
        fmt_v(r_d.iloc[0]["系数"], r_d.iloc[0]["标准误"], r_d.iloc[0]["p值"]) if len(r_d)>0 else "",
        n_val
    ])

make_three_line_table(doc, headers_4a, data_4a, col_widths=[3.5, 2.5, 2.5, 2.5, 2.0])

add_note(doc, '注：所有回归均控制企业FE、年份FE及标准控制变量集。城市层面聚类标准误。')

doc.add_paragraph()

add_para(doc,
    'Dea对供应链集中度、客户集中度、供应商集中度、客户地理距离和供应商地理距离的回归中，'
    '所有系数均不显著（p值在0.07至0.65之间）。唯一的例外是Dea → 供应商集中度1'
    '（系数=-0.145，p=0.073），仅在10%水平上边际显著。'
    'Dea并不直接改变企业可观测的供应链结构。')

# Panel B: 交互效应
add_para(doc, 'Panel B：供应链集中度 × DEA 交互效应', bold=True, size=10, indent_first_line=False)

pc = t4[t4["Panel"] == "C_交互项"]
headers_4c = [["交互项", "系数", "p值", "N"]]
data_4c = []
for _, row in pc.iterrows():
    data_4c.append([
        str(row["解释变量"]),
        fmt_v(row["系数"], row["标准误"], row["p值"]),
        f"{row['p值']:.4f}" if pd.notna(row['p值']) else "",
        f"{int(row['N']):,}" if pd.notna(row['N']) else ""
    ])

make_three_line_table(doc, headers_4c, data_4c, col_widths=[4.0, 3.0, 2.5, 2.5])

doc.add_paragraph()

# Panel C: 分组
add_para(doc, 'Panel C：高/低供应链集中度分组（中位数分割）', bold=True, size=10, indent_first_line=False)

pd4 = t4[t4["Panel"] == "D_分组"]
headers_4d = [["分组", "Dea系数", "p值", "N"]]
data_4d = []
for _, row in pd4.iterrows():
    g = str(row["被解释变量"])
    data_4d.append([
        g,
        fmt_v(row["系数"], row["标准误"], row["p值"]),
        f"{row['p值']:.4f}",
        f"{int(row['N']):,}"
    ])

make_three_line_table(doc, headers_4d, data_4d, col_widths=[5.0, 3.0, 2.5, 2.5])

add_note(doc, '注：按供应链集中度中位数分组。低组=低于中位数，高组=高于中位数。')

doc.add_paragraph()

add_para(doc,
    'Dea × 供应链集中度的交互项高度显著（0.0005，p<0.001）。按中位数分组后：'
    '在供应链集中度较低的组中，Dea不显著（0.0022，p=0.597）；在集中度较高的组中，'
    'Dea显著为正（0.0179，p<0.001），系数约为低组的8倍。')

add_para(doc,
    '这一发现实际上不支持论文的叙事框架（"数据能力降低集中度以增强韧性"），而支持一个替代性解读：'
    '数据要素应用能力的作用在于在高集中度环境下帮助企业更有效地管理供应链关系，'
    '而非改变集中度本身。这是一种补偿效应（compensatory effect），而非结构转变效应'
    '（structural transformation effect）。')

# ═══════════════════════════════════════════════════════════════════════════════
# Section 4: 进一步分析
# ═══════════════════════════════════════════════════════════════════════════════

add_heading_styled(doc, '四、进一步分析：倒U型与被质疑的测量独立性', level=2)

# Inverted-U
add_para(doc, 'Panel A：倒U型关系检验', bold=True, size=10, indent_first_line=False)

iu = t5[t5["Panel"] == "A_倒U型"]
headers_iu = [["变量", "一次项", "二次项", "判定"]]
data_iu = []
for _, row in iu.iterrows():
    data_iu.append([
        row["变量"],
        f"{row['一次项系数']:.4f} ({row['一次项标准误']:.4f})" if pd.notna(row['一次项系数']) else "",
        f"{row['二次项系数']:.4f}**" if pd.notna(row['二次项p']) and row['二次项p'] < 0.05 else f"{row['二次项系数']:.4f}" if pd.notna(row['二次项系数']) else "",
        "显著（正U型）" if pd.notna(row['二次项p']) and row['二次项p'] < 0.05 else "不显著"
    ])

make_three_line_table(doc, headers_iu, data_iu, col_widths=[2.5, 4.0, 4.0, 3.0])

add_note(doc, '注：一次项和二次项均去均值化后进入回归。若二次项显著为正，暗示正U型而非倒U型关系。')

doc.add_paragraph()

add_para(doc,
    'Dea的二次项显著为正（0.0044，p=0.025），暗示Dea与Res之间可能存在正U型而非倒U型关系——'
    '即数据能力的回报在低水平和高水平阶段更为突出。这与论文的理论预期方向相反，需要进一步的理论审视。')

# AI Horse Race
add_para(doc, 'Panel B：AI词频竞争（Horse Race）', bold=True, size=10, indent_first_line=False)

hr = t5[t5["Panel"] == "C_HorseRace"]
headers_hr = [["设定", "Dea系数", "p值(Dea)", "AI_freq系数", "p值(AI)", "判定"]]
data_hr = []
for _, row in hr.iterrows():
    var = str(row["变量"])
    if var == "仅Dea":
        data_hr.append(["仅Dea", "0.0115***", "<0.001", "—", "—", "Dea显著"])
    elif var == "仅AI_freq":
        data_hr.append(["仅AI词频", "—", "—", "0.0081**", "0.050", "AI词频边际显著"])
    elif var == "Dea+AI_freq":
        data_hr.append(["Dea + AI词频", "0.0102***", "<0.001", "0.0039", "0.349", "仅Dea保持显著"])
    elif "corr" in var:
        data_hr.append(["corr(Dea, AI_freq)", "0.737", "", "", "", "高度相关"])

make_three_line_table(doc, headers_hr, data_hr, col_widths=[4.0, 2.5, 2.0, 2.5, 2.0, 3.0])

add_note(doc, '注：AI词频 = ln(1 + 年报AI相关词汇出现次数)。corr(Dea, AI_freq) = 0.737。')

doc.add_paragraph()

add_para(doc,
    '论文核心变量Dea与年报AI词频的相关系数高达0.737。将两者同时纳入回归后，Dea保持显著'
    '（0.0102，p<0.001），而AI词频（0.0039，p=0.349）不再显著。表面上看，这似乎验证了Dea的增量信息。'
    '但0.74的相关性意味着Dea和AI词频共享了超过一半的方差——Dea可能在很大程度上捕捉的是企业'
    '"谈论技术"的通用文本倾向，而非数据要素应用的特定能力。两者的区分效度（discriminant validity）存疑。')

# ═══════════════════════════════════════════════════════════════════════════════
# Section 5: 稳健性检验
# ═══════════════════════════════════════════════════════════════════════════════

add_heading_styled(doc, '五、稳健性检验：通过的项目与致命的失败', level=2)

add_para(doc, '（一）通过的项目', bold=True, size=11, indent_first_line=False)

# Placebo
add_para(doc, '1. 安慰剂检验', bold=True, size=10, indent_first_line=False)

headers_pl = [["变量", "实际系数", "参数p值", "经验p值", "置换均值", "97.5分位", "有效置换数"]]
data_pl = []
for _, row in pl.iterrows():
    data_pl.append([
        row["核心变量"],
        f"{row['实际系数']:.4f}",
        f"{row['参数p值']:.4e}",
        f"{row['经验p值']:.3f}",
        f"{row['置换均值']:.4f}",
        f"{row['97.5分位']:.4f}",
        f"{int(row['有效置换数']):,}"
    ])

make_three_line_table(doc, headers_pl, data_pl, col_widths=[2.5, 2.5, 2.5, 2.5, 2.5, 2.5, 2.5])

add_note(doc, '注：5000次随机置换Dea（Breadth/Depth）的企业归属。经验p值 = 置换系数超过实际系数的比例。')

doc.add_paragraph()

add_para(doc,
    '5000次随机置换后，实际系数0.0115远超置换分布（均值=-0.0006，97.5分位=0.0016），'
    '经验p值=0.000。Dea/Breadth/Depth三个变量均通过安慰剂检验。')

# PSM
add_para(doc, '2. PSM匹配', bold=True, size=10, indent_first_line=False)

headers_psm = [["变量", "原始偏差%", "匹配后偏差%"]]
data_psm = []
for _, row in psm.iterrows():
    data_psm.append([row["变量"], f"{row['原始偏差%']:.1f}", f"{row['匹配后偏差%']:.1f}"])

make_three_line_table(doc, headers_psm, data_psm, col_widths=[4.0, 4.0, 4.0])

add_note(doc, '注：1:1最近邻匹配，卡尺=0.05。匹配后各协变量标准化偏差均小于10%，表明匹配质量良好。')

doc.add_paragraph()

add_para(doc,
    '1:1最近邻匹配后，所有协变量的标准化偏差大幅缩小（如lnrd从67.6%降至6.9%，'
    'lnage从-29.1%降至1.7%），匹配后回归的推断一致性增强。')

# IV
add_para(doc, '（二）致命的失败', bold=True, size=11, indent_first_line=False)

add_para(doc, '1. 工具变量回归', bold=True, size=10, indent_first_line=False)

headers_iv = [["", "系数", "标准误", "p值", "N"]]
data_iv = []
data_iv.append(["OLS (基准)", f"{iv.iloc[0]['OLS_coef']:.4f}", f"({iv.iloc[0]['OLS_se']:.4f})", f"{iv.iloc[0]['OLS_p']:.4e}", f"{int(iv.iloc[0]['OLS_N']):,}"])
data_iv.append(["第一阶段 (雷电→Dea)", f"{iv.iloc[0]['FS_coef']:.4f}", f"({iv.iloc[0]['FS_se']:.4f})", f"{iv.iloc[0]['FS_p']:.4f}", ""])
data_iv.append(["第一阶段F统计量", f"{iv.iloc[0]['FS_F']:.4f}", "", "远低于临界值16.38", ""])
data_iv.append(["2SLS", f"{iv.iloc[0]['2SLS_coef']:.4f}", f"({iv.iloc[0]['2SLS_se']:.4f})", f"{iv.iloc[0]['2SLS_p']:.4e}", f"{int(iv.iloc[0]['2SLS_N']):,}"])

make_three_line_table(doc, headers_iv, data_iv, col_widths=[4.5, 2.5, 2.5, 3.5, 2.0])

add_note(doc, '注：Stock-Yogo弱工具变量检验：10% maximal IV size临界值≈16.38。第一阶段F=0.11，远低于临界值。')

doc.add_paragraph()

add_para(doc,
    '以地级市雷电频率作为Dea的工具变量进行2SLS估计。第一阶段F统计量为0.112，'
    '远低于Stock-Yogo弱工具变量检验的临界值（10% maximal IV size ≈ 16.38）。'
    '这意味着雷电频率几乎完全不能预测Dea，排他性约束的讨论也就失去了前提。IV估计无法为因果识别提供任何额外支撑。')

# Spec curve
add_para(doc, '2. 设定曲线分析', bold=True, size=10, indent_first_line=False)

n_total = len(spec)
n_sig_05 = (spec["pval"] < 0.05).sum()
n_sig_10 = (spec["pval"] < 0.10).sum()

headers_sc = [["指标", "数值"]]
data_sc = [
    ["总设定数", f"{n_total}"],
    ["p<0.05的设定数（比例）", f"{n_sig_05} ({n_sig_05/n_total*100:.1f}%)"],
    ["p<0.10的设定数（比例）", f"{n_sig_10} ({n_sig_10/n_total*100:.1f}%)"],
    ["中位系数", f"{spec['coef'].median():.4f}"],
    ["系数范围", f"[{spec['coef'].min():.4f}, {spec['coef'].max():.4f}]"],
    ["平均p值", f"{spec['pval'].mean():.4f}"],
]

make_three_line_table(doc, headers_sc, data_sc, col_widths=[5.5, 6.5])

add_note(doc, '注：共枚举90种设定组合（FE×聚类方式×控制变量集×样本定义）。城市层面聚类标准误。')

doc.add_paragraph()

add_para(doc,
    f'我们枚举了{n_total}种合理的模型设定组合（3种固定效应 × 2种聚类方式 × 3种控制变量集 × 5种样本定义），'
    f'仅{n_sig_05/n_total*100:.1f}%的设定产生了p<0.05的显著结果，'
    f'{n_sig_10/n_total*100:.1f}%产生p<0.10的结果。中位系数仅为{spec["coef"].median():.4f}，'
    f'系数范围从{spec["coef"].min():.4f}到{spec["coef"].max():.4f}，平均p值为{spec["pval"].mean():.4f}。'
    f'这些数字表明，论文的核心结论高度依赖于特定的设定选择，在不同但同样合理的设定下并不稳健。')

doc.add_paragraph()

add_para(doc, '3. 排除政策混淆', bold=True, size=10, indent_first_line=False)

add_para(doc,
    '控制同期数字政策冲击后，Dea系数变化幅度为0.0%，同期数字政策不构成混淆因素。'
    '该项检验获得通过，但需注意政策变量的构造和样本期间选择可能影响检验的检验力。')

# ═══════════════════════════════════════════════════════════════════════════════
# Section 6: DDD
# ═══════════════════════════════════════════════════════════════════════════════

add_heading_styled(doc, '六、DDD三重差分模型', level=2)

add_para(doc,
    '以数据要素市场化配置政策（2015年起）为外生冲击，引入五个调节变量进行三重差分估计。'
    '模型吸收企业FE、年份-行业FE和年份-省份FE三维固定效应，聚类在企业层面。')

# DDD results (from running 13_ddd_model.py)
ddd_results = [
    ["业务复杂度",   "0.0178*", "(0.0095)", "-0.0230***", "(0.0087)", "34,137"],
    ["供应链集中度", "0.0157*", "(0.0095)", "0.0013**",   "(0.0005)", "34,069"],
    ["本地数商数量", "-0.0068", "(0.0129)", "0.0163",     "(0.0103)", "31,100"],
    ["专利数量",     "0.0142",  "(0.0100)", "-0.0024",    "(0.0058)", "34,747"],
    ["纵向一体化",   "0.0079",  "(0.0095)", "0.0905*",    "(0.0470)", "31,146"],
]

headers_ddd = [["M调节变量", "β₄(DataAsset×Post)", "标准误", "β₇(DataAsset×Post×M)", "标准误", "N"]]
make_three_line_table(doc, headers_ddd, ddd_results, col_widths=[3.5, 3.5, 2.5, 3.5, 2.5, 2.0])

add_note(doc, '注：Y = β₁DataAsset + β₂Post + β₃M + β₄(DataAsset×Post) + β₅(DataAsset×M) + β₆(Post×M) + β₇(DataAsset×Post×M) + ΣControls + FE + ε。')
add_note(doc, '表中仅汇报 β₄ 与 β₇。FE: 企业FE + 年份-行业FE + 年份-省份FE。* p<0.1, ** p<0.05, *** p<0.01。')

doc.add_paragraph()

add_para(doc,
    '政策后数据资产对TFP_LP的正向效应在业务复杂度高（β₄=0.0178，p<0.1）和供应链集中度高的企业中得到了统计支持。'
    '三重交互项（β₇）在业务复杂度（负向，即高复杂度削弱政策效应）和供应链集中度（正向，即高集中度放大政策效应）'
    '两个维度上显著。然而，约一半的β₄和β₇不显著，DDD模型的整体证据强度为中等偏弱。'
    '另外需注意，海外国家数和海外子公司数两项数据不可获取，M变量仅保留了5个可用维度。')

# ═══════════════════════════════════════════════════════════════════════════════
# Section 7: 总体评估
# ═══════════════════════════════════════════════════════════════════════════════

add_heading_styled(doc, '七、总体评估', level=2)

add_para(doc,
    '复现工作确认了论文基准回归的方向和显著性，但也揭示了若干值得警惕的问题：')

add_para(doc,
    '统计层面：（1）Breadth与Depth之间存在高度共线性，使得两者的独立效应不可识别；'
    '（2）所有异质性维度的组间差异均未通过正式检验；'
    '（3）IV估计因弱工具变量而完全失效；'
    '（4）设定曲线分析表明仅有少数设定组合显著。', indent_first_line=True)

add_para(doc,
    '理论层面：（1）DEA与AI词频的高相关性（r=0.74）质疑了其测量特定构念'
    '（数据要素应用能力）的区分效度；'
    '（2）供应链网络分析揭示的是补偿效应而非论文所暗示的结构转变效应。', indent_first_line=True)

add_para(doc,
    '建议：（1）使用文本长度归一化或残差化DEA以排除文本量混淆；'
    '（2）提供DEA与外部数据能力指标（如IT专利、数字资产投资）的相关性证据；'
    '（3）报告所有关键结论的设定曲线以增强透明度；'
    '（4）使用双重机器学习（Chernozhukov et al., 2018）等现代因果推断方法进行稳健性补充。', indent_first_line=True)

# ═══════════════════════════════════════════════════════════════════════════════
# Save
# ═══════════════════════════════════════════════════════════════════════════════

output_path = OUT / "复现结果报告_完整版.docx"
doc.save(str(output_path))
print(f"报告已保存: {output_path}")
print(f"文件大小: {output_path.stat().st_size / 1024:.0f} KB")
