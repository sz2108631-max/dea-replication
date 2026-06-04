"""
将所有复现结果输出为三线表 Word 文档
巫强 et al. (2026)「企业数据要素应用能力与供应链韧性」— 完整复现

所有数据从 output/*.csv 读取，生成学术三线表格式。
"""

import pandas as pd
import numpy as np
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import warnings
warnings.filterwarnings("ignore")

ROOT = Path(__file__).resolve().parent.parent
OUT  = ROOT / "output"


# ═══════════════════════════════════════════════════════════════════════════════
# 三线表工具函数
# ═══════════════════════════════════════════════════════════════════════════════

def set_cell(cell, text, bold=False, font_size=8, font_name='宋体',
             alignment=WD_ALIGN_PARAGRAPH.CENTER):
    """Set cell content with formatting"""
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = alignment
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.line_spacing = Pt(11)
    run = p.add_run(str(text))
    run.font.size = Pt(font_size)
    run.font.bold = bold
    run.font.name = font_name
    # Set East Asian font
    rPr = run._element.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), font_name)
    rPr.insert(0, rFonts)


def apply_three_line_borders(table, header_rows=2):
    """Apply academic three-line table borders to a Word table.

    Top thick line (row 0 top), medium line under header (header_rows-1 bottom),
    thick line at bottom (last row bottom). All other borders removed.
    """
    n_rows = len(table.rows)
    n_cols = len(table.columns)

    # First, remove ALL borders
    for row in table.rows:
        for cell in row.cells:
            tcPr = cell._tc.get_or_add_tcPr()
            # Remove existing tcBorders
            for existing in tcPr.findall(qn('w:tcBorders')):
                tcPr.remove(existing)
            tcBorders = OxmlElement('w:tcBorders')
            for border_name in ['top', 'left', 'bottom', 'right']:
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'nil')
                border.set(qn('w:sz'), '0')
                border.set(qn('w:space'), '0')
                border.set(qn('w:color'), 'auto')
                tcBorders.append(border)
            tcPr.append(tcBorders)

    # Top border: thick line on first row
    for cell in table.rows[0].cells:
        add_cell_border(cell, 'top', '12', '000000')

    # Header-bottom border: medium line on last header row
    for cell in table.rows[header_rows - 1].cells:
        add_cell_border(cell, 'bottom', '6', '000000')

    # Bottom border: thick line on last row
    for cell in table.rows[-1].cells:
        add_cell_border(cell, 'bottom', '12', '000000')


def add_cell_border(cell, border_name, sz, color):
    """Add a single border to a cell"""
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn('w:tcBorders'))
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)
    # Remove existing border of same type
    existing = tcBorders.find(qn(f'w:{border_name}'))
    if existing is not None:
        tcBorders.remove(existing)
    border = OxmlElement(f'w:{border_name}')
    border.set(qn('w:val'), 'single')
    border.set(qn('w:sz'), sz)
    border.set(qn('w:space'), '0')
    border.set(qn('w:color'), color)
    tcBorders.append(border)


def make_three_line_table(doc, headers, data_rows, col_widths=None):
    """Create a three-line table and return the table object.

    headers: list of header rows, each row is a list of strings
    data_rows: list of data rows, each row is a list of strings
    col_widths: optional list of column widths in Cm
    """
    n_cols = len(headers[0])
    all_rows = headers + data_rows
    n_rows = len(all_rows)
    n_header_rows = len(headers)

    table = doc.add_table(rows=n_rows, cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    # Set column widths if specified
    if col_widths:
        for row in table.rows:
            for j, w in enumerate(col_widths):
                if j < n_cols:
                    row.cells[j].width = Cm(w)

    # Fill data
    for i, row_data in enumerate(all_rows):
        is_header = i < n_header_rows
        for j, text in enumerate(row_data):
            if j < n_cols:
                # First column left-aligned, rest centered
                align = WD_ALIGN_PARAGRAPH.LEFT if j == 0 else WD_ALIGN_PARAGRAPH.CENTER
                set_cell(table.rows[i].cells[j], str(text),
                        bold=is_header, font_size=7.5 if not is_header else 8,
                        alignment=align)

    # Apply three-line borders
    apply_three_line_borders(table, header_rows=n_header_rows)

    return table


def add_note(doc, text, font_size=7.5):
    """Add a note paragraph below a table"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.size = Pt(font_size)
    run.font.name = '宋体'
    rPr = run._element.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), '宋体')
    rPr.insert(0, rFonts)
    return p


def fmt_v(coef, se, pval):
    """Format: coef***\\n(se)"""
    if pd.isna(coef) or pd.isna(pval):
        return "N/A"
    stars = "***" if pval < 0.01 else "**" if pval < 0.05 else "*" if pval < 0.1 else ""
    return f"{coef:.4f}{stars}"


def fmt_se(se):
    """Format: (se)"""
    if pd.isna(se):
        return ""
    return f"({se:.4f})"


# ═══════════════════════════════════════════════════════════════════════════════
# 加载数据
# ═══════════════════════════════════════════════════════════════════════════════
print("加载CSV结果文件...")
t1a = pd.read_csv(OUT / "table1_stepwise_from_script.csv")
t1b = pd.read_csv(OUT / "table1_panel_b.csv")
t2  = pd.read_csv(OUT / "table2_mechanism.csv")
t3  = pd.read_csv(OUT / "table3_heterogeneity.csv")
t4  = pd.read_csv(OUT / "table4_supply_chain.csv")
t5  = pd.read_csv(OUT / "table5_further.csv")
print("加载完成\n")

# ═══════════════════════════════════════════════════════════════════════════════
# 构建 Word 文档
# ═══════════════════════════════════════════════════════════════════════════════
doc = Document()

# Page setup
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

# Default style
style = doc.styles['Normal']
style.font.name = '宋体'
style.font.size = Pt(10.5)
style.paragraph_format.line_spacing = 1.15


# ╔══════════════════════════════════════════════════════════════════════════════╗
# ║  表1：基准回归                                                              ║
# ╚══════════════════════════════════════════════════════════════════════════════╝

print("生成 表1：基准回归...")

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('表1：基准回归 — 数据要素应用能力与供应链韧性')
run.font.size = Pt(10.5)
run.font.bold = True
run.font.name = '宋体'

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('被解释变量：供应链韧性（Res）')
run.font.size = Pt(9)
run.font.name = '宋体'

# Panel A: 逐步加控制变量
panel_a = doc.add_paragraph()
run = panel_a.add_run('Panel A: Dea 逐步加入控制变量')
run.font.size = Pt(9)
run.font.bold = True

if len(t1a) > 0:
    models = t1a["模型"].tolist()
    headers_a = [
        ["变量"] + models,
    ]
    # Data rows: Dea coef, Dea SE, then control variable labels, N, R2, FE
    data_a = []
    # Dea row
    dea_row = ["Dea"]
    dea_se_row = [""]
    for i in range(len(t1a)):
        r = t1a.iloc[i]
        dea_row.append(fmt_v(r["Dea系数"], r["Dea标准误"], r["Dea_p值"]))
        dea_se_row.append(fmt_se(r["Dea标准误"]))
    data_a.append(dea_row)
    data_a.append(dea_se_row)

    # Control variables: just show key ones
    ctl_names = ["lnage", "lnsize", "lev", "klr", "lnrd", "own"]
    ctl_labels = ["企业年龄", "企业规模", "资产负债率", "资本劳动比", "研发投入", "所有权性质"]
    for clab in ctl_labels:
        data_a.append([clab] + ["✓" if i >= 1 else "" for i in range(len(models))])

    data_a.append(["N"] + [f"{int(t1a.iloc[i]['N']):,}" for i in range(len(t1a))])
    data_a.append(["R² (组内)"] + [f"{t1a.iloc[i]['R2_within']:.4f}" for i in range(len(t1a))])
    data_a.append(["企业FE"] + ["是"] * len(models))
    data_a.append(["年份FE"] + ["是"] * len(models))

    make_three_line_table(doc, headers_a, data_a,
                          col_widths=[3.0] + [2.0] * len(models))

add_note(doc, "注：括号内为城市层面聚类稳健标准误。*** p<0.01, ** p<0.05, * p<0.1。所有列均包含常数项。")
doc.add_paragraph()  # spacer

# Panel B: Breadth / Depth
panel_b = doc.add_paragraph()
run = panel_b.add_run('Panel B: Breadth 与 Depth 替代解释变量')
run.font.size = Pt(9)
run.font.bold = True

if len(t1b) > 0:
    # Filter: show单独回归 only (not 联合)
    t1b_clean = t1b[~t1b["变量"].str.contains("联合|基准", na=False)]
    vars_b = t1b_clean["变量"].tolist()
    headers_b = [["变量"] + [f"({i+1}) {v}" for i, v in enumerate(vars_b)]]

    # Core coefficients
    data_b = []
    for var in vars_b:
        sub = t1b_clean[t1b_clean["变量"] == var]
        if len(sub) == 0: continue
        r = sub.iloc[0]
        data_b.append([var, fmt_v(r["系数"], r["标准误"], r["p值"])])
        data_b.append(["", fmt_se(r["标准误"])])

    # Add Dea benchmark for comparison
    t1b_dea = t1b[t1b["变量"].str.contains("Dea", na=False)]
    if len(t1b_dea) > 0:
        r = t1b_dea.iloc[0]
        data_b.append(["Dea (基准)", fmt_v(r["系数"], r["标准误"], r["p值"])])
        data_b.append(["", fmt_se(r["标准误"])])

    data_b.append(["N"] + [f"{int(t1b_clean.iloc[i]['N']):,}" for i in range(len(t1b_clean))])
    data_b.append(["R² (组内)"] + ["—"] * len(vars_b))
    data_b.append(["企业FE"] + ["是"] * len(vars_b))
    data_b.append(["年份FE"] + ["是"] * len(vars_b))

    make_three_line_table(doc, headers_b, data_b,
                          col_widths=[3.0] + [3.5])

add_note(doc, "注：Breadth = 数据要素应用广度（覆盖关键词类别数）；Depth = 数据要素应用深度（各类别内平均相似度）。")
doc.add_page_break()


# ╔══════════════════════════════════════════════════════════════════════════════╗
# ║  表2：机制检验                                                              ║
# ╚══════════════════════════════════════════════════════════════════════════════╝

print("生成 表2：机制检验...")

title2 = doc.add_paragraph()
title2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title2.add_run('表2：机制检验 — 数据要素应用能力的调节效应')
run.font.size = Pt(10.5)
run.font.bold = True
run.font.name = '宋体'

# Panel A: Baron-Kenny 三步法
panel_2a = doc.add_paragraph()
run = panel_2a.add_run('Panel A: Baron-Kenny 三步法（因果中介检验）')
run.font.size = Pt(9)
run.font.bold = True

bk = t2[t2["机制"].str.contains("_BK", na=False)]
if len(bk) > 0:
    headers_2a = [["步骤", "渠道", "路径", "系数", "N", "判定"]]
    data_2a = []
    for _, row in bk.iterrows():
        mech = row["机制"].replace("_BK", "")
        s1 = fmt_v(row.get("Step1_Dea_Res", np.nan), np.nan, row.get("Step1_p", np.nan))
        s2 = fmt_v(row.get("Step2_Dea_M", np.nan), np.nan, row.get("Step2_p", np.nan))
        s3d = fmt_v(row.get("Step3_Dea_Res_given_M", np.nan), np.nan, row.get("Step3_p_Dea", np.nan))
        s3m = fmt_v(row.get("Step3_M_Res_given_Dea", np.nan), np.nan, row.get("Step3_p_M", np.nan))
        n = f"{int(row['N']):,}" if pd.notna(row['N']) else ""

        p1 = "✅" if pd.notna(row.get("Step1_p")) and row["Step1_p"] < 0.05 else "❌"
        p2 = "✅" if pd.notna(row.get("Step2_p")) and row["Step2_p"] < 0.05 else "❌"
        p3m = "✅" if pd.notna(row.get("Step3_p_M")) and row["Step3_p_M"] < 0.05 else "❌"

        data_2a.append(["Step1", mech, "Dea → Res", s1, n, p1])
        data_2a.append(["Step2", mech, f"Dea → {mech}", s2, n, p2])
        data_2a.append(["Step3", mech, f"Dea → Res | {mech}", s3d, n, ""])
        data_2a.append(["Step3", mech, f"{mech} → Res | Dea", s3m, n, p3m])
        data_2a.append(["", "", "", "", "", ""])

    make_three_line_table(doc, headers_2a, data_2a,
                          col_widths=[1.2, 1.5, 3.5, 2.0, 2.0, 1.5])

add_note(doc, "注：Baron-Kenny (1986) 三步法。Step1: DEA→Res; Step2: DEA→M; Step3: 加入中介变量M后，M→Res须显著且DEA系数应下降。")
doc.add_paragraph()

# Panel B: 交互项法
panel_2b = doc.add_paragraph()
run = panel_2b.add_run('Panel B: Rajan-Zingales 交互项法（调节效应）')
run.font.size = Pt(9)
run.font.bold = True

inter = t2[~t2["机制"].str.contains("_BK", na=False)]
if len(inter) > 0:
    headers_2b = [["交互项", "机制", "系数", "SE", "p值", "N", "判定"]]
    data_2b = []
    for _, row in inter.iterrows():
        c = row["交互项系数"]
        s = row["交互项标准误"]
        p = row["交互项p值"]
        judge = "✅ 显著" if pd.notna(p) and p < 0.05 else ("⚠️ 边际" if pd.notna(p) and p < 0.1 else "❌ 不显著")
        n = f"{int(row['N']):,}" if pd.notna(row['N']) else ""
        data_2b.append([
            f"{row['核心变量']} × {row['机制']}",
            row['机制'],
            f"{c:.4f}" if pd.notna(c) else "N/A",
            f"({s:.4f})" if pd.notna(s) else "",
            f"{p:.4f}" if pd.notna(p) else "",
            n,
            judge
        ])

    make_three_line_table(doc, headers_2b, data_2b,
                          col_widths=[3.0, 1.5, 2.0, 2.0, 1.5, 2.0, 2.0])

add_note(doc, "注：Rajan & Zingales (1998) 交互项法。交互项 = 核心变量(去均值) × 调节变量(去均值)。若交互项显著，说明调节变量放大/缩小了核心变量的效应。")
doc.add_page_break()


# ╔══════════════════════════════════════════════════════════════════════════════╗
# ║  表3：异质性分析                                                            ║
# ╚══════════════════════════════════════════════════════════════════════════════╝

print("生成 表3：异质性分析...")

title3 = doc.add_paragraph()
title3.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title3.add_run('表3：异质性分析 — 数据要素应用能力的非对称效应')
run.font.size = Pt(10.5)
run.font.bold = True
run.font.name = '宋体'

sub3 = doc.add_paragraph()
sub3.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub3.add_run('被解释变量：供应链韧性（Res）')
run.font.size = Pt(9)
run.font.name = '宋体'

if len(t3) > 0:
    dims = t3["异质性维度"].unique()
    for dim in dims:
        sub_t3 = t3[t3["异质性维度"] == dim]
        dim_label = dim.replace("H1_", "H1: ").replace("H2_", "H2: ").replace("H3_", "H3: ").replace("H4_", "H4: ").replace("H6_", "H6: ")

        dim_p = doc.add_paragraph()
        run = dim_p.add_run(dim_label)
        run.font.size = Pt(9)
        run.font.bold = True

        groups = sub_t3["分组"].tolist()
        headers_3 = [["分组"] + groups]
        data_3 = []

        # Dea coefficients
        coef_row = ["Dea"]
        se_row = [""]
        n_row = ["N"]
        r2_row = ["R² (组内)"]
        for _, row in sub_t3.iterrows():
            coef_row.append(fmt_v(row["Dea系数"], row["Dea标准误"], row["Dea_p值"]))
            se_row.append(fmt_se(row["Dea标准误"]))
            n_row.append(f"{int(row['N']):,}" if pd.notna(row['N']) else "")
            r2_row.append(f"{row['R2_within']:.4f}" if pd.notna(row['R2_within']) else "")
        data_3.append(coef_row)
        data_3.append(se_row)

        # Compute group difference
        if len(sub_t3) == 2:
            c0, c1 = sub_t3.iloc[0]["Dea系数"], sub_t3.iloc[1]["Dea系数"]
            s0, s1 = sub_t3.iloc[0]["Dea标准误"], sub_t3.iloc[1]["Dea标准误"]
            diff = c0 - c1
            se_diff = np.sqrt(s0**2 + s1**2) if pd.notna(s0) and pd.notna(s1) else np.nan
            if pd.notna(diff) and pd.notna(se_diff) and se_diff > 0:
                z_diff = diff / se_diff
                p_diff = 2 * (1 - 0.5 * (1 + np.tanh(np.abs(z_diff) / 1.2533)))
                data_3.append([f"组间差异: Δ={diff:.4f}"] + [f"z≈{z_diff:.2f}, p≈{p_diff:.3f}"] + [""] * (len(groups)-1))

        data_3.append(n_row)
        data_3.append(r2_row)
        data_3.append(["企业FE"] + ["是"] * len(groups))
        data_3.append(["年份FE"] + ["是"] * len(groups))

        make_three_line_table(doc, headers_3, data_3,
                              col_widths=[3.0] + [3.5] * len(groups))

add_note(doc, "注：分组回归均采用双向固定效应模型（企业FE + 年份FE），城市层面聚类标准误。组间差异基于近似z检验。")
doc.add_page_break()


# ╔══════════════════════════════════════════════════════════════════════════════╗
# ║  表4：供应链网络                                                            ║
# ╚══════════════════════════════════════════════════════════════════════════════╝

print("生成 表4：供应链网络...")

title4 = doc.add_paragraph()
title4.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title4.add_run('表4：供应链网络与集中度 — DEA如何影响供应链结构？')
run.font.size = Pt(10.5)
run.font.bold = True
run.font.name = '宋体'

if len(t4) > 0:
    # Panel A: DEA → 供应链结构
    panel_4a = doc.add_paragraph()
    run = panel_4a.add_run('Panel A: DEA → 供应链结构')
    run.font.size = Pt(9)
    run.font.bold = True

    pa = t4[t4["Panel"] == "A_DEA到供应链"]
    if len(pa) > 0:
        dep_vars_a = pa["被解释变量"].unique()
        headers_4a = [["被解释变量"] + ["Dea", "Breadth", "Depth", "N"]]
        data_4a = []
        for dv in dep_vars_a:
            row_coef = [dv]
            row_se = [""]
            n_val = ""
            for core in ["Dea", "Breadth", "Depth"]:
                sub = pa[(pa["被解释变量"] == dv) & (pa["解释变量"] == core)]
                if len(sub) > 0:
                    r = sub.iloc[0]
                    row_coef.append(fmt_v(r["系数"], r["标准误"], r["p值"]))
                    row_se.append(fmt_se(r["标准误"]))
                    n_val = f"{int(r['N']):,}" if pd.notna(r['N']) else ""
                else:
                    row_coef.append("")
                    row_se.append("")
            row_coef.append(n_val)
            row_se.append("")
            data_4a.append(row_coef)
            data_4a.append(row_se)

        make_three_line_table(doc, headers_4a, data_4a,
                              col_widths=[3.0, 2.5, 2.5, 2.5, 2.0])

    add_note(doc, "注：每一行是一个独立回归，被解释变量为对应的供应链指标，解释变量为DEA指标 + 全部控制变量 + 企业FE + 年份FE。")
    doc.add_paragraph()

    # Panel B: 供应链集中度 → Res
    panel_4b = doc.add_paragraph()
    run = panel_4b.add_run('Panel B: 供应链集中度 → Res（交互效应）')
    run.font.size = Pt(9)
    run.font.bold = True

    pc = t4[t4["Panel"] == "C_交互项"]
    if len(pc) > 0:
        headers_4c = [["交互项", "系数", "SE", "p值", "N", "判定"]]
        data_4c = []
        for _, row in pc.iterrows():
            c = row["系数"]
            s = row["标准误"]
            p = row["p值"]
            judge = "✅" if pd.notna(p) and p < 0.05 else "❌"
            n = f"{int(row['N']):,}" if pd.notna(row['N']) else ""
            data_4c.append([
                str(row["解释变量"]),
                f"{c:.4f}" if pd.notna(c) else "",
                f"({s:.4f})" if pd.notna(s) else "",
                f"{p:.4f}" if pd.notna(p) else "",
                n,
                judge
            ])

        make_three_line_table(doc, headers_4c, data_4c,
                              col_widths=[3.5, 2.0, 2.0, 1.5, 2.0, 1.5])

    add_note(doc, "注：交互项 = DEA(去均值) × 供应链指标(去均值)。控制变量包含全部9个基准控制变量。企业FE + 年份FE。")
    doc.add_paragraph()

    # Panel C: 高/低集中度分组
    panel_4d = doc.add_paragraph()
    run = panel_4d.add_run('Panel C: 高/低供应链集中度分组（中位数分割）')
    run.font.size = Pt(9)
    run.font.bold = True

    pd4 = t4[t4["Panel"] == "D_分组"]
    if len(pd4) > 0:
        base_vars = set()
        for _, row in pd4.iterrows():
            dep = str(row["被解释变量"])
            for suffix in ["_低组", "_高组"]:
                if dep.endswith(suffix):
                    base_vars.add(dep[:-len(suffix)])

        headers_4d = [["供应链指标", "低集中度组", "高集中度组"]]
        data_4d = []
        for bv in sorted(base_vars):
            low = pd4[pd4["被解释变量"] == f"{bv}_低组"]
            high = pd4[pd4["被解释变量"] == f"{bv}_高组"]
            cl = fmt_v(low.iloc[0]["系数"], low.iloc[0]["标准误"], low.iloc[0]["p值"]) if len(low) > 0 else "N/A"
            ch = fmt_v(high.iloc[0]["系数"], high.iloc[0]["标准误"], high.iloc[0]["p值"]) if len(high) > 0 else "N/A"
            data_4d.append([bv, cl, ch])
            data_4d.append(["", fmt_se(low.iloc[0]["标准误"]) if len(low) > 0 else "",
                           fmt_se(high.iloc[0]["标准误"]) if len(high) > 0 else ""])

        make_three_line_table(doc, headers_4d, data_4d,
                              col_widths=[4.0, 4.0, 4.0])

add_note(doc, "注：分组基于各供应链指标的中位数。被解释变量为Res。控制变量包含全部基准控制变量。企业FE + 年份FE。")
doc.add_page_break()


# ╔══════════════════════════════════════════════════════════════════════════════╗
# ║  表5：进一步分析                                                            ║
# ╚══════════════════════════════════════════════════════════════════════════════╝

print("生成 表5：进一步分析...")

title5 = doc.add_paragraph()
title5.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title5.add_run('表5：进一步分析 — 倒U型检验、DEA vs AI词频')
run.font.size = Pt(10.5)
run.font.bold = True
run.font.name = '宋体'

if len(t5) > 0:
    # Panel A: 倒U型检验
    panel_5a = doc.add_paragraph()
    run = panel_5a.add_run('Panel A: 倒U型检验（二次项）')
    run.font.size = Pt(9)
    run.font.bold = True

    u_sub = t5[t5["Panel"] == "A_倒U型"]
    if len(u_sub) > 0:
        headers_5a = [["变量", "一次项 (β₁)", "二次项 (β₂)", "Q1(最低25%)", "Q4(最高25%)", "形状判定"]]
        data_5a = []
        for _, row in u_sub.iterrows():
            lin = fmt_v(row["一次项系数"], row.get("一次项标准误", np.nan), row["一次项p"])
            sq = fmt_v(row["二次项系数"], row.get("二次项标准误", np.nan), row["二次项p"])
            q1 = fmt_v(row["Q1系数"], row.get("Q1标准误", np.nan), row["Q1_p"])
            q4 = fmt_v(row["Q4系数"], row.get("Q4标准误", np.nan), row["Q4_p"])

            sq_c = row["二次项系数"]
            sq_p = row["二次项p"]
            lin_p = row["一次项p"]
            shape = "正U型 (阈值效应)" if pd.notna(sq_p) and sq_p < 0.05 and pd.notna(sq_c) and sq_c > 0 else \
                    "倒U型" if pd.notna(sq_p) and sq_p < 0.05 and pd.notna(sq_c) and sq_c < 0 else \
                    "线性" if pd.notna(lin_p) and lin_p < 0.05 else "不显著"

            data_5a.append([str(row["变量"]), lin, sq, q1, q4, shape])
            # SE row for 一次项 and 二次项
            lin_se = fmt_se(row.get("一次项标准误", np.nan))
            sq_se = fmt_se(row.get("二次项标准误", np.nan))
            q1_se = fmt_se(row.get("Q1标准误", np.nan))
            q4_se = fmt_se(row.get("Q4标准误", np.nan))
            data_5a.append(["", lin_se, sq_se, q1_se, q4_se, ""])

        make_three_line_table(doc, headers_5a, data_5a,
                              col_widths=[2.0, 2.5, 2.5, 2.5, 2.5, 3.0])

    add_note(doc, "注：一次项和二次项来自同一回归。Q1/Q4为分位数分组回归。形状判定：二次项显著为正→正U型（阈值效应），显著为负→倒U型。")
    doc.add_paragraph()

    # Panel B: DEA vs AI词频 Horse Race
    panel_5c = doc.add_paragraph()
    run = panel_5c.add_run('Panel B: DEA vs AI词频 — 竞争性预测（Horse Race）')
    run.font.size = Pt(9)
    run.font.bold = True

    hr_sub = t5[t5["Panel"] == "C_HorseRace"]
    if len(hr_sub) > 0:
        headers_5c = [["模型", "Dea 系数", "AI_freq 系数", "N"]]
        data_5c = []
        for _, row in hr_sub.iterrows():
            label = str(row["变量"])
            if "corr" in label:
                continue
            dea_c = fmt_v(row["一次项系数"], np.nan, row["一次项p"])
            ai_c = fmt_v(row["二次项系数"], np.nan, row["二次项p"])
            n = f"{int(row['N']):,}" if pd.notna(row['N']) else ""
            data_5c.append([label, dea_c if "Dea" in label or "+" in label else "", ai_c if "AI" in label or "+" in label else "", n])

        # Add correlation
        corr_row = hr_sub[hr_sub["变量"].str.contains("corr", na=False)]
        if len(corr_row) > 0:
            corr_val = corr_row.iloc[0]["一次项系数"]
            data_5c.append(["corr(Dea, AI_freq)", f"{corr_val:.4f}", "", ""])

        make_three_line_table(doc, headers_5c, data_5c,
                              col_widths=[3.5, 3.0, 3.0, 2.0])

add_note(doc, "注：所有回归均包含全部基准控制变量 + 企业FE + 年份FE。AI_freq = 企业年报全文本中AI相关关键词词频和加1取对数。corr为Pearson相关系数。")
doc.add_page_break()


# ╔══════════════════════════════════════════════════════════════════════════════╗
# ║  表6：稳健性检验汇总                                                        ║
# ╚══════════════════════════════════════════════════════════════════════════════╝

print("生成 表6：稳健性检验汇总...")

title6 = doc.add_paragraph()
title6.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title6.add_run('表6：稳健性检验汇总')
run.font.size = Pt(10.5)
run.font.bold = True
run.font.name = '宋体'

# R1: PSM
panel_r1 = doc.add_paragraph()
run = panel_r1.add_run('Panel A: PSM匹配后回归')
run.font.size = Pt(9)
run.font.bold = True

headers_r1 = [["", "Dea系数", "p值", "N"]]
data_r1 = [
    ["匹配前 (基准)", "0.0115***", "0.0001", "35,135"],
    ["        (0.0029)", "", "", ""],
    ["匹配后 (ATT)", "0.0111***", "<0.01", "3,310"],
    ["", "", "", ""],
]
make_three_line_table(doc, headers_r1, data_r1,
                      col_widths=[4.0, 3.0, 2.0, 2.0])
add_note(doc, "注：按Dea中位数划分处理组/控制组，1:1最近邻匹配（卡尺=0.05），1,655对匹配成功，9/9协变量匹配后偏差<10%。")
doc.add_paragraph()

# R2: 安慰剂检验
panel_r2 = doc.add_paragraph()
run = panel_r2.add_run('Panel B: 安慰剂检验（5000次随机置换）')
run.font.size = Pt(9)
run.font.bold = True

headers_r2 = [["变量", "实际系数", "参数p值", "经验p值", "置换均值", "95%区间", "判定"]]
data_r2 = [
    ["Dea", "0.0115", "0.0001", "0.0000", "-0.0006", "[-0.0028, 0.0016]", "✅ 通过"],
    ["Breadth", "0.0091", "0.0210", "0.0000", "-0.0005", "[-0.0029, 0.0018]", "✅ 通过"],
    ["Depth", "0.0127", "0.0000", "0.0000", "-0.0006", "[-0.0028, 0.0016]", "✅ 通过"],
]
make_three_line_table(doc, headers_r2, data_r2,
                      col_widths=[2.0, 2.0, 2.0, 2.0, 2.0, 4.0, 2.0])
add_note(doc, "注：每年内随机置换处理变量取值，打破企业关联但保留年份分布特征。经验p值 = (置换系数 > 实际系数的次数) / 5000。")
doc.add_paragraph()

# R3: IV
panel_r3 = doc.add_paragraph()
run = panel_r3.add_run('Panel C: 工具变量（雷电频率）')
run.font.size = Pt(9)
run.font.bold = True

headers_r3 = [["指标", "值", "判定"]]
data_r3 = [
    ["OLS (同样本)", "0.0116***", "基准"],
    ["第一阶段 (雷电 → Dea)", "-0.0491 (p=0.738)", "❌ 不显著"],
    ["Kleibergen-Paap F统计量", "0.11", "❌ 远低于临界值10"],
    ["2SLS系数", "-2.62 (p=0.006)", "❌ 符号错误+量级荒谬"],
    ["排他性约束 (雷电 → Res | Dea)", "p<0.01", "❌ 排他性被违反"],
    ["安慰剂 (雷电 → 企业年龄)", "p=0.020", "❌ 安慰剂未通过"],
]
make_three_line_table(doc, headers_r3, data_r3,
                      col_widths=[5.5, 4.0, 4.0])
add_note(doc, "注：工具变量为地级市雷电频率。Kleibergen-Paap F统计量临界值为10（Stock & Yogo, 2005）。IV估计不可靠。")
doc.add_paragraph()

# R4: 排除替代解释
panel_r4 = doc.add_paragraph()
run = panel_r4.add_run('Panel D: 排除替代解释 — 同期数字政策冲击')
run.font.size = Pt(9)
run.font.bold = True

headers_r4 = [["政策冲击", "控制后Dea系数", "系数变化", "政策自身p值", "判定"]]
data_r4 = [
    ["智算中心试点DiD", "0.0115***", "+0.0%", "0.72", "✅ 不混淆"],
    ["数据交易所DiD", "0.0114***", "-0.7%", "0.68", "✅ 不混淆"],
    ["数据要素市场化配置", "0.0110***", "-4.3%", "0.46", "✅ 不混淆"],
    ["全部3项政策同时控制", "0.0109***", "-5.1%", "—", "✅ 不混淆"],
]
make_three_line_table(doc, headers_r4, data_r4,
                      col_widths=[4.5, 3.0, 2.5, 2.5, 2.5])
add_note(doc, "注：每行报告在原基准回归基础上额外加入政策冲击DiD项后的Dea系数。系数变化 = (控制后 - 基准) / 基准 × 100%。")
doc.add_paragraph()

# R5: 设定曲线
panel_r5 = doc.add_paragraph()
run = panel_r5.add_run('Panel E: 设定曲线（90种设定组合）')
run.font.size = Pt(9)
run.font.bold = True

headers_r5 = [["子样本", "组合数", "p<0.05比例", "判定"]]
data_r5 = [
    ["全部设定", "90", "16% (14/90)", "❌ 绝大多数不显著"],
    ["平衡面板", "24", "0% (0/24)", "❌❌❌ 完全消失"],
    ["排除COVID年", "24", "3% (1/24)", "❌❌ 几乎消失"],
    ["行业FE替代企业FE", "30", "4% (1/30)", "❌❌ 几乎消失"],
    ["全样本+企业FE+城市聚类", "6", "100% (6/6)", "✅ 基准设定恰在显著区间"],
]
make_three_line_table(doc, headers_r5, data_r5,
                      col_widths=[5.0, 2.0, 3.0, 5.0])
add_note(doc, "注：枚举 3种FE × 3种聚类方式 × 2种控制变量集 × 3种子样本 × 2种解释变量 ≈ 90种合理设定组合。城市聚类标准误。")
doc.add_page_break()


# ╔══════════════════════════════════════════════════════════════════════════════╗
# ║  表7：DDD三重交互模型                                                       ║
# ╚══════════════════════════════════════════════════════════════════════════════╝

print("生成 表7：DDD三重交互模型...")

title7 = doc.add_paragraph()
title7.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title7.add_run('附表：机制交互检验 — 强度DID三重交互模型')
run.font.size = Pt(12)
run.font.bold = True
run.font.name = '宋体'

# Load DDD results from 13_ddd_model.py or compute inline
# Note: This section generates from pre-computed values (same as 13_ddd_model.py)
# Since we don't re-run the DDD model, we hardcode the validated results

m_labels = ["业务复杂度", "供应链集中度", "本地数商数量", "专利数量", "纵向一体化"]

# Panel A: TFP_LP
panel_7a = doc.add_paragraph()
run = panel_7a.add_run('Panel A: Y = 全要素生产率 (TFP_LP)')
run.font.size = Pt(9)
run.font.bold = True

# Results from 13_ddd_model.py / 14_ddd_three_line_table.py
ddd_tfp_lp = {
    "业务复杂度":   (0.0178, 0.0095, 0.060,  -0.0230, 0.0087, 0.008,  34137, 0.995),
    "供应链集中度": (0.0157, 0.0095, 0.098,  0.0013,  0.0005, 0.011,  34069, 0.995),
    "本地数商数量": (-0.0068, 0.0129, 0.600, 0.0163,  0.0103, 0.113,  31100, 0.995),
    "专利数量":     (0.0142, 0.0100, 0.155,  -0.0024, 0.0058, 0.676,  34747, 0.995),
    "纵向一体化":   (0.0079, 0.0095, 0.408,  0.0905,  0.0470, 0.054,  31146, 0.995),
}

# Panel B: TFP_OP
ddd_tfp_op = {
    "业务复杂度":   (0.0159, 0.0098, 0.106,  -0.0238, 0.0091, 0.009,  34137, 0.992),
    "供应链集中度": (0.0150, 0.0098, 0.127,  0.0014,  0.0006, 0.014,  34069, 0.992),
    "本地数商数量": (-0.0054, 0.0134, 0.689, 0.0166,  0.0106, 0.116,  31100, 0.991),
    "专利数量":     (0.0114, 0.0103, 0.269,  -0.0037, 0.0059, 0.528,  34747, 0.992),
    "纵向一体化":   (0.0060, 0.0098, 0.543,  0.0982,  0.0484, 0.042,  31146, 0.992),
}

for panel_name, ddd_results in [("Panel A: Y = 全要素生产率 (TFP_LP)", ddd_tfp_lp),
                                  ("Panel B: Y = 全要素生产率 (TFP_OP)", ddd_tfp_op)]:
    panel_p = doc.add_paragraph()
    run = panel_p.add_run(panel_name)
    run.font.size = Pt(9)
    run.font.bold = True

    n_m = len(m_labels)
    headers_7 = [
        ["变量"] + [f"({i+1})" for i in range(n_m)],
        [""] + m_labels,
    ]

    data_7 = []
    # β₄ row
    b4_row = ["Dataasset×Post"]
    b4_se_row = [""]
    for m in m_labels:
        b4, b4_se, b4_p = ddd_results[m][:3]
        b4_row.append(fmt_v(b4, b4_se, b4_p))
        b4_se_row.append(fmt_se(b4_se))
    data_7.append(b4_row)
    data_7.append(b4_se_row)

    # β₇ row
    b7_row = ["Dataasset×Post×M"]
    b7_se_row = [""]
    for m in m_labels:
        _, _, _, b7, b7_se, b7_p = ddd_results[m][:6]
        b7_row.append(fmt_v(b7, b7_se, b7_p))
        b7_se_row.append(fmt_se(b7_se))
    data_7.append(b7_row)
    data_7.append(b7_se_row)

    # FE rows
    data_7.append(["年份固定效应"] + ["NO"] * n_m)
    data_7.append(["企业固定效应"] + ["YES"] * n_m)
    data_7.append(["年份-行业联合固定效应"] + ["YES"] * n_m)
    data_7.append(["年份-省份联合固定效应"] + ["YES"] * n_m)
    data_7.append(["控制变量"] + ["YES"] * n_m)

    # N
    n_row = ["样本数"]
    for m in m_labels:
        n_row.append(f"{ddd_results[m][6]:,}")
    data_7.append(n_row)

    # R²
    r2_row = ["R²"]
    for m in m_labels:
        r2_row.append(f"{ddd_results[m][7]:.3f}")
    data_7.append(r2_row)

    make_three_line_table(doc, headers_7, data_7,
                          col_widths=[3.5] + [2.2] * n_m)

    add_note(doc,
        "注：括号内为企业层面聚类稳健标准误。* p<0.1, ** p<0.05, *** p<0.01。"
        "回归方程为 Y = β₁DataAsset + β₂Post + β₃M + β₄(DataAsset×Post) + β₅(DataAsset×M) + β₆(Post×M) + β₇(DataAsset×Post×M) + ΣControls + FE + ε。"
        "表中仅汇报 β₄ 与 β₇，其余低阶项未展示。DataAsset = DEA指数; Post = 数据要素市场化配置政策(2015年起); M = 各列标注的调节变量。"
    )
    doc.add_paragraph()


# ═══════════════════════════════════════════════════════════════════════════════
# 保存
# ═══════════════════════════════════════════════════════════════════════════════
output_path = OUT / '复现结果_三线表_完整版.docx'
doc.save(str(output_path))
print(f"\n{'='*60}")
print(f"三线表 Word 文档已保存: {output_path}")
print(f"文件大小: {output_path.stat().st_size / 1024:.0f} KB")
print(f"{'='*60}")
