"""
生成原论文与复现数据的描述性统计对比表 → 单个 Word 文件
"""
import pyreadstat
import pandas as pd
import numpy as np
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parent.parent
OUT = ROOT / "output"
DATA_ORIG = "/Users/weixuan/Documents/论文复现/CMDA_管理层讨论与分析_ALL/原文及附件/数据.dta"
DATA_OWN = OUT / "replication_panel_own.dta"

# 统一变量定义：变量名 → (中文标签, 顺序)
VARIABLES = [
    ('res', '供应链韧性'),
    ('Dea', '数据要素应用能力'),
    ('Breadth', '应用广度'),
    ('Depth', '应用深度'),
    ('Dea_count', 'DEA关键词计数'),
    ('Breadth_count', '广度关键词计数'),
    ('Depth_count', '深度关键词计数'),
    ('lnage', 'ln(企业年龄)'),
    ('lnsize', 'ln(企业规模)'),
    ('klr', '资本劳动比'),
    ('lev', '杠杆率'),
    ('debt', '负债率'),
    ('lnrd', 'ln(研发投入)'),
    ('bsize', 'ln(董事会规模)'),
    ('dual', '两职合一'),
    ('indrate', '独立董事比例'),
    ('own', '股权集中度'),
    ('IC', '内部控制指数'),
    ('cost', '交易成本'),
]


def compute_descriptive(df, var_list):
    """计算描述性统计"""
    rows = []
    for var, label in var_list:
        if var not in df.columns:
            continue
        s = df[var].dropna()
        if len(s) == 0:
            continue
        rows.append({
            '变量': label,
            '样本数': len(s),
            '均值': s.mean(),
            '标准差': s.std(),
            '最小值': s.min(),
            'P25': s.quantile(0.25),
            '中位数': s.median(),
            'P75': s.quantile(0.75),
            '最大值': s.max(),
        })
    return pd.DataFrame(rows)


def format_num(val, fmt='.3f'):
    """格式化数字"""
    if pd.isna(val):
        return ''
    abs_val = abs(val)
    if abs_val >= 1000:
        return f'{val:,.2f}'
    elif abs_val >= 10:
        return f'{val:{fmt}}'
    elif abs_val >= 1:
        return f'{val:.4f}'
    elif abs_val >= 0.01:
        return f'{val:.4f}'
    else:
        return f'{val:.4f}'


def set_cell(cell, text, bold=False, font_size=9, align='center'):
    """设置单元格文本"""
    cell.text = ''
    p = cell.paragraphs[0]
    if align == 'center':
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == 'left':
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(str(text))
    run.font.size = Pt(font_size)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.bold = bold


def add_stats_table(doc, title, stats_df, note=None):
    """向 Word 添加描述性统计表"""
    # 标题
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    cols = list(stats_df.columns)
    nrows = len(stats_df)
    ncols = len(cols)

    table = doc.add_table(rows=nrows + 1, cols=ncols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    # 表头
    for j, col in enumerate(cols):
        set_cell(table.cell(0, j), col, bold=True, font_size=8)

    # 数据
    for i, (_, row) in enumerate(stats_df.iterrows()):
        for j, col in enumerate(cols):
            val = row[col]
            if col == '变量':
                set_cell(table.cell(i + 1, j), str(val), font_size=8, align='left')
            elif col == '样本数':
                set_cell(table.cell(i + 1, j), f'{int(val):,}', font_size=8)
            else:
                set_cell(table.cell(i + 1, j), format_num(val), font_size=8)

    # 注释
    if note:
        p_note = doc.add_paragraph()
        run_note = p_note.add_run(note)
        run_note.font.size = Pt(7)
        run_note.font.name = '宋体'
        run_note._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    doc.add_paragraph()


def main():
    # 加载数据
    print("加载数据...")
    df_orig, _ = pyreadstat.read_dta(DATA_ORIG)
    df_own, _ = pyreadstat.read_dta(DATA_OWN)
    print(f"  原论文: {len(df_orig):,} obs")
    print(f"  复现:   {len(df_own):,} obs")

    # 筛选可用变量
    var_list_orig = [(v, l) for v, l in VARIABLES if v in df_orig.columns]
    var_list_own = [(v, l) for v, l in VARIABLES if v in df_own.columns]

    # 计算描述性统计
    print("计算描述性统计...")
    stats_orig = compute_descriptive(df_orig, var_list_orig)
    stats_own = compute_descriptive(df_own, var_list_own)

    # 合并对比表 — 统一变量名，取两表交集变量
    common_vars = []
    for v, label in VARIABLES:
        in_orig = v in df_orig.columns
        in_own = v in df_own.columns
        if in_orig or in_own:
            common_vars.append((v, label, in_orig, in_own))

    # Panel A: 原论文
    print(f"  原论文: {len(stats_orig)} 变量")
    print(f"  复现:   {len(stats_own)} 变量")

    # 创建 Word
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)

    # 大标题
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run('描述性统计\n巫强等 (2026) 原论文 vs 复现数据')
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    doc.add_paragraph()

    # 数据概览
    overview = doc.add_paragraph()
    run = overview.add_run(
        f'原论文样本: {len(df_orig):,} 观测值 | 复现样本: {len(df_own):,} 观测值\n'
        f'原论文变量数: {len(df_orig.columns)} | 复现变量数: {len(df_own.columns)}'
    )
    run.font.size = Pt(9)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    doc.add_paragraph()

    # 原论文描述性统计
    add_stats_table(doc, 'Panel A: 原论文数据描述性统计', stats_orig,
                    note='数据来源：巫强等(2026)原始数据（数据.dta）。')

    # 复现数据描述性统计
    add_stats_table(doc, 'Panel B: 复现数据描述性统计（自测算DEA）', stats_own,
                    note='数据来源：自测算DEA + CSMAR财务数据 + 迪博IC + 供应链地理距离 + PageRank + 雷电频率等。')

    # 变量说明
    doc.add_paragraph()
    note_p = doc.add_paragraph()
    note_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = note_p.add_run('变量说明：')
    run.bold = True
    run.font.size = Pt(8)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    notes_text = [
        'res = 供应链韧性（论文构造的综合指标）',
        'Dea = 数据要素应用能力（自测算/原文通过BERT+MD&A文本分析构造）',
        'Breadth = 应用广度（关键词覆盖的维度数量）',
        'Depth = 应用深度（关键词在各维度的深入程度）',
        'lnage = ln(企业成立年限+1)',
        'lnsize = ln(总资产)',
        'klr = 资本劳动比 = 固定资产/员工人数',
        'lev = 资产负债率',
        'debt = 负债率（仅原论文有）',
        'lnrd = ln(研发支出+1)，缺失值未做均值填充',
        'bsize = ln(董事会人数)',
        'dual = 两职合一（董事长兼任总经理=1）',
        'indrate = 独立董事占比',
        'own = 第一大股东持股比例',
        'IC = 迪博内部控制指数',
        'cost = 交易成本（营业费用/营业收入）',
    ]
    for note in notes_text:
        p = doc.add_paragraph()
        run = p.add_run(note)
        run.font.size = Pt(7)
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # 保存
    path = OUT / 'descriptive_statistics.docx'
    doc.save(path)
    print(f"\n输出: {path}")


if __name__ == '__main__':
    main()
